from flask import Flask, render_template, request, redirect, url_for,flash,session,send_file, make_response
from flask_mysqldb import MySQL
from docx import Document
from datetime import datetime
import smtplib
import io

app = Flask(__name__)
app.secret_key = 'your_secret_key'

# Konfigurasi MySQL
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = ''
app.config['MYSQL_DB'] = 'sekolah'
mysql = MySQL(app)

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/form')
def form():
    return render_template('form.html')

@app.route('/submit', methods=['POST'])
def submit():
    if request.method == 'POST':
        nama_pendaftar = request.form['nama_pendaftar']
        tanggal_lahir = request.form['tanggal_lahir']
        alamat_email = request.form['alamat_email']
        nomor_aktif = request.form['nomor_aktif']
        alamat_rumah = request.form['alamat_rumah']
        detail_ortu = request.form['detail_ortu']
        waktu_daftar = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        cur = mysql.connection.cursor()
        cur.execute("INSERT INTO pendaftaran (nama_pendaftar, tanggal_lahir, alamat_email, nomor_aktif, alamat_rumah, detail_ortu, waktu_daftar) VALUES (%s, %s, %s, %s, %s, %s, %s)", 
                    (nama_pendaftar, tanggal_lahir, alamat_email, nomor_aktif, alamat_rumah, detail_ortu, waktu_daftar))
        mysql.connection.commit()
        cur.close()

        return redirect(url_for('result', nama_pendaftar=nama_pendaftar, tanggal_lahir=tanggal_lahir, alamat_email=alamat_email, nomor_aktif=nomor_aktif, alamat_rumah=alamat_rumah, detail_ortu=detail_ortu, waktu_daftar=waktu_daftar))

@app.route('/result')
def result():
    nama_pendaftar = request.args.get('nama_pendaftar')
    tanggal_lahir = request.args.get('tanggal_lahir')
    alamat_email = request.args.get('alamat_email')
    nomor_aktif = request.args.get('nomor_aktif')
    alamat_rumah = request.args.get('alamat_rumah')
    detail_ortu = request.args.get('detail_ortu')
    waktu_daftar = request.args.get('waktu_daftar')
    return render_template('result.html', nama_pendaftar=nama_pendaftar, tanggal_lahir=tanggal_lahir, alamat_email=alamat_email, nomor_aktif=nomor_aktif, alamat_rumah=alamat_rumah, detail_ortu=detail_ortu, waktu_daftar=waktu_daftar)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        if username == 'admin' and password == 'admin123':  # Ubah sesuai dengan username dan password admin Anda
            session['admin'] = True
            return redirect(url_for('admin_dashboard'))
        else:
            flash('Login gagal. Periksa username dan password Anda.', 'danger')
    
    return render_template('login.html')

@app.route('/admin_dashboard')
def admin_dashboard():
    if not session.get('admin'):
        return redirect(url_for('login'))

    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM pendaftaran")
    pendaftar = cur.fetchall()
    cur.close()

    return render_template('admin_dashboard.html', pendaftar=pendaftar)

@app.route('/send_message/<int:id>', methods=['POST'])
def send_message(id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    cur = mysql.connection.cursor()
    cur.execute("SELECT alamat_email FROM pendaftaran WHERE id = %s", [id])
    user = cur.fetchone()
    email = user[0]
    cur.close()

    # Kirim pesan melalui email
    send_email(email, "Pendaftaran Diterima", "Anda telah diterima di sekolah ini 😊")

    flash('Pesan telah dikirim.', 'success')
    return redirect(url_for('admin_dashboard'))

def send_email(to_email, subject, body):
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login('sdn01buton@gmail.com', 'butonraya123')
        message = f'Subject: {subject}\n\n{body}'
        server.sendmail('snd01buton@gmail.com', to_email, message)
        server.quit()
    except Exception as e:
        print(f'Error sending email: {e}')

@app.route('/logout')
def logout():
    session.pop('admin', None)
    return redirect(url_for('home'))

@app.route('/download_word')
def download_word():
    nama_pendaftar = request.args.get('nama_pendaftar')
    tanggal_lahir = request.args.get('tanggal_lahir')
    alamat_email = request.args.get('alamat_email')
    nomor_aktif = request.args.get('nomor_aktif')
    alamat_rumah = request.args.get('alamat_rumah')
    detail_ortu = request.args.get('detail_ortu')
    waktu_daftar = request.args.get('waktu_daftar')

    doc = Document()
    doc.add_heading('Data Pendaftaran', 0)

    doc.add_paragraph(f'Nama Pendaftar: {nama_pendaftar}')
    doc.add_paragraph(f'Tanggal Lahir: {tanggal_lahir}')
    doc.add_paragraph(f'Alamat Email: {alamat_email}')
    doc.add_paragraph(f'Nomor Aktif/WA: {nomor_aktif}')
    doc.add_paragraph(f'Alamat Rumah: {alamat_rumah}')
    doc.add_paragraph(f'Detail Ortu/Wali: {detail_ortu}')
    doc.add_paragraph(f'Waktu Daftar: {waktu_daftar}')

    # Simpan dokumen ke buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Mengirimkan file sebagai respons unduhan
    response = make_response(send_file(buffer, as_attachment=True, download_name=f"pendaftaran_{nama_pendaftar.replace(' ', '_')}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
    response.headers['Content-Disposition'] = f'attachment; filename=pendaftaran_{nama_pendaftar.replace(" ", "_")}.docx'
    
    return response

if __name__ == '__main__':
    app.run(debug=True)
